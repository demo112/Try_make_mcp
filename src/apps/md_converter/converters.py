import os
import logging
import markdown
from bs4 import BeautifulSoup
from openpyxl import Workbook
from docx import Document
from xhtml2pdf import pisa
import platform

# 离线安全回调：禁止网络请求
def link_callback(uri, rel):
    """
    Convert HTML URIs to absolute system paths so xhtml2pdf can access those resources
    Enforces OFFLINE ONLY rule: rejects http/https/ftp
    """
    # 1. 拒绝网络协议
    if uri.startswith('http://') or uri.startswith('https://') or uri.startswith('ftp://'):
        logging.warning(f"Offline Policy: Blocked network resource: {uri}")
        return None

    # 2. 处理本地文件
    # 如果是绝对路径，直接返回
    if os.path.isabs(uri):
        if os.path.exists(uri):
            return uri
        return None
        
    # 3. 处理相对路径 (需要知道 base path，这里简化处理，假设相对于 cwd)
    # 实际上 pisa.CreatePDF 应该通过 context 传递 base path，但在 MCP 中通常文件是完全限定的
    # 如果 rel 有定义，尝试结合
    # 暂时只支持绝对路径的本地资源，或者与 cwd 相关的
    abs_path = os.path.abspath(uri)
    if os.path.exists(abs_path):
        return abs_path
        
    return None

# 设置 Windows 字体路径
# 优先使用系统中确实存在的字体，避免临时文件权限问题
WINDOWS_FONT_PATH = "C:/Windows/Fonts/simhei.ttf"  # 黑体
# 如果黑体不存在，尝试使用微软雅黑
if not os.path.exists(WINDOWS_FONT_PATH):
    WINDOWS_FONT_PATH = "C:/Windows/Fonts/msyh.ttc"

def _get_html_with_font(html_content: str) -> str:
    """
    为 HTML 添加中文字体支持的 CSS。
    注意：xhtml2pdf 在处理本地字体路径时，如果直接使用 Windows 路径可能会有问题。
    一种常见的做法是将字体文件读取并转换为 base64，但这会增加代码量。
    另一种是使用相对路径或 file:// 协议，但 Windows 下也有兼容性问题。
    最简单的修复尝试：不使用 @font-face 加载本地路径，而是依赖 xhtml2pdf 内部机制（虽然它默认不带中文字体）。
    或者，尝试将 font_path 转换为 file uri。
    """
    # 尝试修复：使用 reportlab 注册字体（xhtml2pdf 基于 reportlab）
    # 但为了保持代码简洁，我们先尝试调整 font-face 的写法。
    
    # 确保字体文件存在，如果不存在则不添加自定义字体 CSS，避免报错
    if not os.path.exists(WINDOWS_FONT_PATH):
        return f"<html><head><meta charset='utf-8'></head><body>{html_content}</body></html>"

    # 修复：在 Windows 上，reportlab/xhtml2pdf 对路径的处理可能需要 file:/// 前缀
    # 或者我们直接加载字体内容（为了稳健性，这里尝试注册字体的方式比较复杂，我们先试着忽略字体错误）
    
    # 备选方案：直接使用系统默认字体，虽然中文会乱码，但至少能生成 PDF。
    # 为了解决中文乱码，必须加载字体。
    
    # 尝试方案 C：手动加载字体到 reportlab
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        pdfmetrics.registerFont(TTFont('CustomFont', WINDOWS_FONT_PATH))
        font_css = "font-family: 'CustomFont';"
    except Exception as e:
        # 如果注册失败，回退到无字体
        print(f"Warning: Failed to register font: {e}")
        font_css = ""

    css = f"""
    <style>
        body {{
            {font_css}
        }}
        p {{
            {font_css}
        }}
        div {{
            {font_css}
        }}
        table {{
            {font_css}
        }}
    </style>
    """
    return f"<html><head><meta charset='utf-8'>{css}</head><body>{html_content}</body></html>"

def md_to_word(source_path: str, output_path: str) -> str:
    """
    Markdown 转 Word (docx)。
    MVP 实现: 简单解析 HTML 并写入 docx (目前仅支持纯文本和简单结构，复杂样式可能丢失)。
    为了更好的效果，通常建议使用 pandoc，但为了纯 Python 依赖，这里做基础实现。
    """
    if not os.path.exists(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")
        
    with open(source_path, "r", encoding="utf-8") as f:
        text = f.read()
        
    # 解析 Markdown 为 HTML
    html = markdown.markdown(text)
    soup = BeautifulSoup(html, "html.parser")
    
    doc = Document()
    
    for element in soup.recursiveChildGenerator():
        if element.name:
            if element.name.startswith("h") and len(element.name) == 2:
                try:
                    level = int(element.name[1])
                    doc.add_heading(element.get_text(), level=level)
                except ValueError:
                    doc.add_paragraph(element.get_text())
            elif element.name == "p":
                doc.add_paragraph(element.get_text())
            elif element.name == "li":
                doc.add_paragraph(element.get_text(), style="List Bullet")
            # 这里可以扩展更多标签支持
            
    # 如果内容为空，至少写入原始内容
    if not soup.get_text().strip():
        doc.add_paragraph(text)
        
    doc.save(output_path)
    return f"Successfully converted to Word: {output_path}"

def md_to_pdf(source_path: str, output_path: str) -> str:
    """
    Markdown 转 PDF。
    """
    if not os.path.exists(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")

    with open(source_path, "r", encoding="utf-8") as f:
        text = f.read()
        
    # 解析 Markdown 为 HTML
    html_body = markdown.markdown(text, extensions=['tables'])
    
    # 包装 HTML 并添加字体 CSS
    full_html = _get_html_with_font(html_body)
    
    with open(output_path, "wb") as result_file:
        pisa_status = pisa.CreatePDF(
            src=full_html,
            dest=result_file,
            encoding='utf-8',
            link_callback=link_callback  # 注入离线回调
        )
        
    if pisa_status.err:
        raise RuntimeError(f"PDF generation error: {pisa_status.err}")
        
    return f"Successfully converted to PDF: {output_path}"

def md_to_excel(source_path: str, output_path: str) -> str:
    """
    Markdown 转 Excel (提取表格)。
    """
    if not os.path.exists(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")

    with open(source_path, "r", encoding="utf-8") as f:
        text = f.read()
        
    # 解析 Markdown 为 HTML (启用表格扩展)
    html = markdown.markdown(text, extensions=['tables'])
    soup = BeautifulSoup(html, "html.parser")
    
    tables = soup.find_all("table")
    if not tables:
        return "No tables found in Markdown file. No Excel file generated."
        
    wb = Workbook()
    # 移除默认创建的 Sheet
    wb.remove(wb.active)
    
    for i, table in enumerate(tables):
        ws = wb.create_sheet(title=f"Table {i+1}")
        
        # 处理表头
        thead = table.find("thead")
        row_idx = 1
        if thead:
            for tr in thead.find_all("tr"):
                col_idx = 1
                for th in tr.find_all("th"):
                    ws.cell(row=row_idx, column=col_idx, value=th.get_text())
                    col_idx += 1
                row_idx += 1
        
        # 处理表体
        tbody = table.find("tbody")
        if tbody:
            for tr in tbody.find_all("tr"):
                col_idx = 1
                for td in tr.find_all("td"):
                    ws.cell(row=row_idx, column=col_idx, value=td.get_text())
                    col_idx += 1
                row_idx += 1
                
        # 如果没有 thead/tbody，尝试直接处理 tr
        if not thead and not tbody:
             for tr in table.find_all("tr"):
                col_idx = 1
                for cell in tr.find_all(["td", "th"]):
                    ws.cell(row=row_idx, column=col_idx, value=cell.get_text())
                    col_idx += 1
                row_idx += 1
                
    wb.save(output_path)
    return f"Successfully extracted {len(tables)} tables to Excel: {output_path}"
